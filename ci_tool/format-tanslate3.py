# -*- coding: utf-8 -*-
import docx
import os
import sys

# 需替换文档目录
# path = u'C:\George\DOCX'
path = 'D:\docs'
# 自动创建


if sys.getdefaultencoding() != 'utf-8':
    reload(sys)
    sys.setdefaultencoding('utf-8')


# 两个日志
def logerr(text):
    errlog = path + u'\替换出错列表.txt'
    with open(errlog, "a+") as f:
        f.write(text)
        f.write('\n')


def logok(text):
    infolog = path + u'\替换成功列表.txt'
    with open(infolog, "a+") as f:
        f.write(text)
        f.write('\n')


# 替换内容（文档名称，旧的内容，新的内容）
def info_update(doc, old_info, new_info):
    # 替换文档中所有文字内容
    for para in doc.paragraphs:
        for run in para.runs:
            run.text = run.text.replace(old_info, new_info)
    # 替换文档中表格中的内容
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = cell.text.replace(old_info, new_info)
    # 替换页眉页脚
    doc.sections[0].header.paragraphs[0].text = doc.sections[0].header.paragraphs[0].text.replace(old_info, new_info)
    doc.sections[0].footer.paragraphs[0].text = doc.sections[0].footer.paragraphs[0].text.replace(old_info, new_info)


def thr(old_info, new_info):
    # 遍历目录中的docx文档
    for parent, dirnames, filenames in os.walk(path):
        for fn in filenames:
            filedir = os.path.join(parent, fn)
            if fn.endswith('.docx'):
                try:
                    # 定义文档路径
                    doc = docx.Document(filedir)
                    # 调用函数修改文档内容
                    info_update(doc, old_info, new_info)
                    # 保存文档
                    doc.save(filedir)
                    # 写入修改日志
                    logok(filedir)
                    print(filedir + ' 修改完成')
                except Exception as e:
                    # 写入修改失败日志
                    logerr(filedir)
                    print(filedir + ' 修改失败')


if __name__ == '__main__':
    if len(sys.argv) != 4:
        for arg in sys.argv:
            print(arg)
        print("参数错误，路径 老词 新词")
        sys.exit(2)
    path = sys.argv[1]
    oldword = sys.argv[2]
    newword = sys.argv[3]
    # oldword = '川建'
    # newword = '川普'
    print(oldword, newword)
    thr(oldword, newword)
    print('----全部替换完成----')