# -*- coding: utf-8 -*-
import docx
import os
import sys

# 需替换文档目录
path = u'C:\George\DOCX'
# 自动创建
tlog = path + u'\替换文档列表.txt'
err_log = path + u'\替换出错列表.txt'

if sys.getdefaultencoding() != 'utf-8':
    reload(sys)
    sys.setdefaultencoding('utf-8')


# 两个日志
def log(text):
    with open(err_log, "a+") as f:
        f.write(text)
        f.write('\n')


def log2(text):
    with open(tlog, "a+") as f:
        f.write(text)
        f.write('\n')


# 替换内容（文档名称，旧的内容，新的内容）
def info_update(doc, old_info, new_info):
    # 替换文档中所有文字内容
    for para in doc.paragraphs:
        for run in para.runs:
            run.text = run.text.replace(old_info, new_info)
    # 替换文档中表格中的内容
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = cell.text.replace(old_info, new_info)


def thr(old_info, new_info):
    # 遍历目录中的docx文档
    for parent, dirnames, filenames in os.walk(path):
        for fn in filenames:
            filedir = os.path.join(parent, fn)
            if fn.endswith('.docx'):
                try:
                    # 定义文档路径
                    doc = docx.Document(filedir)
                    # 调用函数修改文档内容
                    info_update(doc, old_info, new_info)
                    # 保存文档
                    doc.save(filedir)
                    # 写入修改日志
                    log2(filedir + ' 修改完成')
                    print(filedir + ' 修改完成')
                except Exception as e:
                    # 写入修改失败日志
                    log(filedir)


if __name__ == '__main__':
    if len(sys.argv) != 3:
        for arg in sys.argv:
            print(arg)
        print("参数错误，老词 新词")
        sys.exit(2)
    oldword = sys.argv[1]
    newword = sys.argv[2]
    print(oldword, newword)
    thr(oldword, newword)
    print('----全部替换完成----')

